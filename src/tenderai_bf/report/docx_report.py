"""DOCX report generation using python-docx."""

import io
from datetime import datetime
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Cm, Inches, Pt, RGBColor

from ..config import settings
from ..logging import get_logger

logger = get_logger(__name__)


def _add_formatted_text(paragraph, text: str):
    """Add text with markdown-style formatting to a paragraph.
    
    Supports:
    - **bold**
    - *italic*
    - Line breaks
    """
    import re
    
    # Split by markdown bold/italic patterns
    # Pattern: **text** or *text*
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    
    for part in parts:
        if not part:
            continue
        
        # Bold: **text**
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        # Italic: *text*
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
        # Regular text
        else:
            # Replace line breaks
            lines = part.split('\n')
            for i, line in enumerate(lines):
                paragraph.add_run(line)
                if i < len(lines) - 1:
                    paragraph.add_run('\n')


def add_hyperlink(paragraph, text: str, url: str):
    """Add a hyperlink to a paragraph."""
    
    # Create hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), url)
    
    # Create run element
    run = OxmlElement('w:r')
    run_props = OxmlElement('w:rPr')
    
    # Add hyperlink styling
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    run_props.append(color)
    
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    run_props.append(underline)
    
    run.append(run_props)
    
    # Add text
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    run.append(text_elem)
    
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_bookmark(paragraph, bookmark_id: str, bookmark_name: str):
    """Add a bookmark to a paragraph for internal linking."""
    p = paragraph._p
    
    # Add bookmark start
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), str(bookmark_id))
    bookmark_start.set(qn('w:name'), bookmark_name)
    p.insert(0, bookmark_start)
    
    # Add bookmark end
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), str(bookmark_id))
    p.append(bookmark_end)


def add_internal_hyperlink(paragraph, text: str, bookmark_name: str):
    """Add an internal hyperlink to a bookmark within the document."""
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)
    
    # Create run element
    run = OxmlElement('w:r')
    run_props = OxmlElement('w:rPr')
    
    # Add hyperlink styling (blue and underlined)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    run_props.append(color)
    
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    run_props.append(underline)
    
    run.append(run_props)
    
    # Add text
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    run.append(text_elem)
    
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def build_report(data: Dict[str, Any]) -> Optional[bytes]:
    """Build a DOCX report from pipeline data."""
    
    try:
        # Create document
        document = Document()
        
        # Set document properties
        document.core_properties.title = "TenderAI – YULCOM Technologies"
        document.core_properties.author = "TenderAI"
        document.core_properties.subject = "Rapport de veille des appels d'offres IT/Ingénierie"
        document.core_properties.created = data['generated_at']
        
        # Build report sections
        _add_title_page(document, data)
        _add_table_of_contents(document, data)
        _add_executive_summary(document, data)
        _add_notices_section(document, data)
        _add_other_notices_section(document, data)
        _add_sources_section(document, data)
        _add_appendices(document, data)
        
        # Convert to bytes
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        
        report_bytes = buffer.getvalue()
        buffer.close()
        
        logger.info(
            "DOCX report generated successfully",
            size_bytes=len(report_bytes),
            notices_count=len(data.get('notices', [])),
            run_id=data.get('run_id')
        )
        
        return report_bytes
    
    except Exception as e:
        logger.error(
            "Failed to generate DOCX report",
            error=str(e),
            run_id=data.get('run_id'),
            exc_info=True
        )
        return None


def _add_title_page(document: Document, data: Dict[str, Any]) -> None:
    """Add title page to the document."""
    
    # Title
    title = document.add_heading('', level=0)
    title_run = title.runs[0] if title.runs else title.add_run()
    title_run.text = "TenderAI – YULCOM Technologies"
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = document.add_paragraph()
    subtitle_run = subtitle.add_run("Rapport de veille des appels d'offres IT/Ingénierie")
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Date and run info
    document.add_paragraph()  # Empty line
    
    date_para = document.add_paragraph()
    date_run = date_para.add_run(f"Généré le : {data['generated_at'].strftime('%d/%m/%Y à %H:%M UTC')}")
    date_run.font.size = Pt(12)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run_para = document.add_paragraph()
    run_run = run_para.add_run(f"ID d'exécution : {data['run_id']}")
    run_run.font.size = Pt(10)
    run_run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    run_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add some space
    document.add_paragraph()
    document.add_paragraph()
    
    # YULCOM branding
    brand_para = document.add_paragraph()
    brand_run = brand_para.add_run("YULCOM Technologies")
    brand_run.font.size = Pt(14)
    brand_run.font.bold = True
    brand_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    brand_desc = document.add_paragraph()
    brand_desc_run = brand_desc.add_run("Système autonome de veille des appels d'offres")
    brand_desc_run.font.size = Pt(10)
    brand_desc_run.font.italic = True
    brand_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Page break
    document.add_page_break()


def _add_executive_summary(document: Document, data: Dict[str, Any]) -> None:
    """Add executive summary section."""
    
    # Section title with bookmark
    heading = document.add_heading('I. Résumé exécutif', level=1)
    add_bookmark(heading, 0, 'section_I')
    
    # Statistics
    stats = data.get('statistics', {})
    notices = data.get('notices', [])
    sources = data.get('sources', [])
    errors = data.get('errors', [])
    
    # Overview paragraph
    overview = document.add_paragraph()
    overview.add_run(
        f"Ce rapport présente les résultats de la veille automatique des appels d'offres "
        f"IT/Ingénierie au Burkina Faso pour la période du {data['generated_at'].strftime('%d/%m/%Y')}. "
        f"Le système a analysé {len(sources)} sources et identifié {len(notices)} avis pertinents."
    )
    
    # Key metrics table
    document.add_paragraph()
    document.add_heading('Métriques clés', level=2)
    
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Métrique'
    header_cells[1].text = 'Valeur'
    
    # Make header bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Add data rows - excluding reports_generated, emails_sent, and total_time 
    # as they are not available at report generation time
    metrics = [
        ('Sources consultées', str(stats.get('sources_checked', 0))),
        ('Liens découverts', str(stats.get('links_discovered', 0))),
        ('Items analysés', str(stats.get('items_parsed', 0))),
        ('Avis pertinents', str(stats.get('relevant_items', 0))),
        ('Avis uniques (après dédoublonnage)', str(stats.get('unique_items', 0))),
    ]
    
    for metric, value in metrics:
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = value
    
    # Errors summary
    if errors:
        document.add_paragraph()
        document.add_heading('Incidents et erreurs', level=2)
        error_para = document.add_paragraph()
        error_para.add_run(f"⚠️ {len(errors)} erreur(s) rencontrée(s) durant l'exécution. Voir la section Annexes pour plus de détails.")
        error_para.runs[0].font.color.rgb = RGBColor(0xFF, 0x66, 0x00)
    
    document.add_page_break()


def _add_table_of_contents(document: Document, data: Dict[str, Any]) -> None:
    """Add table of contents with clickable internal links."""
    
    document.add_heading('Table des matières', level=1)
    
    toc_items = [
        ("I. Résumé exécutif", "section_I"),
        ("II. Avis d'appels d'offres pertinents", "section_II"),
        ("III. Autres Avis", "section_III"),
        ("IV. Sources consultées", "section_IV"),
        ("V. Annexes", "section_V"),
    ]
    
    # Main sections with internal hyperlinks
    for item, bookmark in toc_items:
        para = document.add_paragraph()
        para.style = 'Normal'
        add_internal_hyperlink(para, item, bookmark)
    
    # Subsections (not clickable, just for reference)
    document.add_paragraph("   V.1. Journal des erreurs", style='Normal')
    document.add_paragraph("   V.2. Statistiques détaillées", style='Normal')
    
    document.add_page_break()


def _add_notices_section(document: Document, data: Dict[str, Any]) -> None:
    """Add notices section with appel_offres only."""
    
    notices = data.get('notices', [])
    
    # Section heading with bookmark
    heading = document.add_heading('II. Avis d\'appels d\'offres pertinents', level=1)
    add_bookmark(heading, 1, 'section_II')
    
    if not notices:
        no_notices = document.add_paragraph()
        no_notices.add_run("Aucun avis d'appel d'offres pertinent trouvé pour cette période.")
        no_notices.runs[0].font.italic = True
        return
    
    # Count appel_offres
    appel_offres_count = sum(1 for n in notices if n.get('type', 'appel_offres') == 'appel_offres')
    
    overview_para = document.add_paragraph()
    overview_para.add_run(
        f"Cette section présente les {appel_offres_count} appel(s) d'offre(s) identifié(s) comme pertinent(s) "
        f"pour les domaines IT/Ingénierie au Burkina Faso."
    )
    
    # Group notices by type
    from collections import defaultdict
    notices_by_type = defaultdict(list)
    for notice in notices:
        notice_type = notice.get('type', 'appel_offres')
        notices_by_type[notice_type].append(notice)
    
    # Display only appel_offres in this section
    notice_index = 1
    
    if 'appel_offres' in notices_by_type:
        for notice in notices_by_type['appel_offres']:
            _add_notice_card(document, notice, notice_index)
            notice_index += 1
    
    document.add_page_break()


def _add_other_notices_section(document: Document, data: Dict[str, Any]) -> None:
    """Add other notices section (Rectificatifs, Prorogations, Communiqués, Annulations, Autres)."""
    
    notices = data.get('notices', [])
    
    # Group notices by type
    from collections import defaultdict
    notices_by_type = defaultdict(list)
    for notice in notices:
        notice_type = notice.get('type', 'appel_offres')
        if notice_type != 'appel_offres':  # Only non-appel_offres
            notices_by_type[notice_type].append(notice)
    
    # If no other notices, add a message and return
    if not notices_by_type:
        heading = document.add_heading('III. Autres Avis', level=1)
        add_bookmark(heading, 2, 'section_III')
        no_notices = document.add_paragraph()
        no_notices.add_run("Aucun autre avis (rectificatif, prorogation, communiqué, annulation) trouvé pour cette période.")
        no_notices.runs[0].font.italic = True
        document.add_page_break()
        return
    
    # Count other notices
    other_count = sum(len(notices) for notices in notices_by_type.values())
    
    # Section heading with bookmark
    heading = document.add_heading('III. Autres Avis', level=1)
    add_bookmark(heading, 2, 'section_III')
    
    overview_para = document.add_paragraph()
    overview_para.add_run(
        f"Cette section présente les {other_count} autre(s) avis (rectificatif, prorogation, communiqué, annulation) "
        f"identifié(s) comme pertinent(s) pour les domaines IT/Ingénierie au Burkina Faso."
    )
    
    # Type labels in French
    type_labels = {
        'rectificatif': 'Rectificatifs',
        'prorogation': 'Prorogations',
        'communique': 'Communiqués',
        'annulation': 'Annulations',
        'autre': 'Autres'
    }
    
    # Display in order: rectificatif, prorogation, communique, annulation, autre
    notice_index = 1
    
    for notice_type in ['rectificatif', 'prorogation', 'communique', 'annulation', 'autre']:
        if notice_type in notices_by_type:
            document.add_paragraph()
            section_label = type_labels.get(notice_type, notice_type.title())
            document.add_heading(section_label, level=2)
            
            for notice in notices_by_type[notice_type]:
                _add_notice_card(document, notice, notice_index)
                notice_index += 1
    
    document.add_page_break()


def _add_notice_card(document: Document, notice: Dict[str, Any], index: int) -> None:
    """Add a single notice card."""
    
    # Add notice heading with title or tender_object
    notice_title = notice.get('tender_object', notice.get('title', 'Titre non disponible'))
    document.add_heading(f"{index}. {notice_title}", level=2)
    
    # Create info table
    table = document.add_table(rows=8, cols=2)
    table.style = 'Table Grid'
    
    # Populate table
    info_items = [
        ('Référence', notice.get('reference', notice.get('ref_no', 'N/A'))),
        ('Entité', notice.get('entity', 'N/A')),
        ('Catégorie', notice.get('category', 'N/A')),
        ('Localisation', notice.get('location', 'N/A')),
        ('Date de publication', notice.get('published_at', 'N/A')),
        ('Date limite', notice.get('deadline', notice.get('deadline_at', 'N/A'))),
    ]
    
    for i, (label, value) in enumerate(info_items):
        row = table.rows[i]
        label_cell = row.cells[0]
        value_cell = row.cells[1]
        
        # Make label bold
        label_para = label_cell.paragraphs[0]
        label_run = label_para.add_run(label)
        label_run.font.bold = True
        
        # Add value
        value_cell.text = str(value)
    
    # Add summary
    document.add_paragraph()
    document.add_heading('Résumé', level=3)
    
    summary_text = notice.get('description', notice.get('summary_fr', 'Résumé non disponible'))
    
    # Parse markdown-style formatting in summary
    summary_para = document.add_paragraph()
    _add_formatted_text(summary_para, summary_text)
    
    # Add source URL
    document.add_paragraph()
    url_para = document.add_paragraph()
    url_para.add_run("Source : ")
    
    # Try source_url first (for PDF quotidiens), then url (for web scraping)
    url = notice.get('source_url') or notice.get('url', '')
    if url:
        try:
            add_hyperlink(url_para, url, url)
        except Exception:
            url_para.add_run(url)
    else:
        url_para.add_run("URL non disponible")
    
    # Add relevance score
    relevance_score = notice.get('relevance_score', 0)
    if relevance_score:
        score_para = document.add_paragraph()
        score_run = score_para.add_run(f"Score de pertinence : {relevance_score:.2f}")
        score_run.font.size = Pt(9)
        score_run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    
    # Add separator (use simple hyphens instead of unicode box drawing)
    document.add_paragraph("-" * 80)
    document.add_paragraph()


def _add_sources_section(document: Document, data: Dict[str, Any]) -> None:
    """Add sources section."""
    
    sources = data.get('sources', [])
    
    # Section heading with bookmark
    heading = document.add_heading('IV. Sources consultées', level=1)
    add_bookmark(heading, 3, 'section_IV')
    
    if not sources:
        document.add_paragraph("Aucune source consultée.")
        return
    
    overview_para = document.add_paragraph()
    overview_para.add_run(
        "Cette section liste les sources de données consultées lors de cette exécution."
    )
    
    # Create sources table
    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # Header row
    header_cells = table.rows[0].cells
    headers = ['Nom', 'Type', 'URL', 'Statut']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Add data rows
    for source in sources:
        row_cells = table.add_row().cells
        row_cells[0].text = source.get('name', 'N/A')
        row_cells[1].text = source.get('parser_type', 'N/A')
        row_cells[2].text = source.get('list_url', 'N/A')
        
        # Status based on last_success_at
        if source.get('last_success_at'):
            status = "✅ Succès"
        elif source.get('last_error_at'):
            status = "❌ Erreur"
        else:
            status = "❓ Inconnu"
        row_cells[3].text = status
    
    document.add_page_break()


def _add_appendices(document: Document, data: Dict[str, Any]) -> None:
    """Add appendices section."""
    
    # Section heading with bookmark
    heading = document.add_heading('V. Annexes', level=1)
    add_bookmark(heading, 4, 'section_V')
    
    # Error log
    errors = data.get('errors', [])
    document.add_heading('V.1. Journal des erreurs', level=2)
    
    if not errors:
        document.add_paragraph("Aucune erreur signalée.")
    else:
        for i, error in enumerate(errors, 1):
            error_para = document.add_paragraph()
            error_para.add_run(f"{i}. [{error.get('step', 'unknown')}] {error.get('error', 'Unknown error')}")
            
            timestamp = error.get('timestamp', '')
            if timestamp:
                time_para = document.add_paragraph()
                time_run = time_para.add_run(f"   Heure : {timestamp}")
                time_run.font.size = Pt(9)
                time_run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    
    # Detailed statistics
    document.add_paragraph()
    document.add_heading('V.2. Statistiques détaillées', level=2)
    
    stats = data.get('statistics', {})
    stats_table = document.add_table(rows=1, cols=2)
    stats_table.style = 'Table Grid'
    
    # Header
    header_cells = stats_table.rows[0].cells
    header_cells[0].text = 'Métrique'
    header_cells[1].text = 'Valeur'
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Add all stats with proper formatting
    for key, value in stats.items():
        row_cells = stats_table.add_row().cells
        row_cells[0].text = key.replace('_', ' ').title()
        
        # Format time values (keys ending with '_seconds')
        if key.endswith('_seconds') and isinstance(value, (int, float)):
            if value < 0.01:
                # Very small values in milliseconds
                row_cells[1].text = f"{value * 1000:.2f} ms"
            elif value < 60:
                # Seconds
                row_cells[1].text = f"{value:.2f} s"
            elif value < 3600:
                # Minutes
                minutes = int(value // 60)
                seconds = value % 60
                row_cells[1].text = f"{minutes} min {seconds:.1f} s"
            else:
                # Hours
                hours = int(value // 3600)
                minutes = int((value % 3600) // 60)
                row_cells[1].text = f"{hours} h {minutes} min"
        else:
            row_cells[1].text = str(value)
    
    # Footer
    document.add_paragraph()
    document.add_paragraph()
    
    footer_para = document.add_paragraph()
    footer_run = footer_para.add_run(
        f"Rapport généré par TenderAI BF v{settings.app_version} - "
        f"YULCOM Technologies - {data['generated_at'].strftime('%d/%m/%Y %H:%M UTC')}"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER